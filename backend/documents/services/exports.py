import hashlib
import io
import logging
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.signing import BadSignature, SignatureExpired, TimestampSigner
from django.db import transaction
from django.utils import timezone
from docx import Document as DocxDocument
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from documents.exceptions import (
    DocumentStateConflict,
    InvalidDownloadToken,
)
from documents.models import (
    DocumentAuditEvent,
    DocumentExport,
    GeneratedDocument,
)
from subscriptions.constants import Feature
from subscriptions.models import UsageRecord
from subscriptions.services.entitlements import authorize_feature
from subscriptions.services.usage import consume_usage

DOWNLOAD_SALT = 'documents.export.download'
logger = logging.getLogger(__name__)


def _latest_revision(document):
    return document.revisions.order_by('-version').first()


@transaction.atomic
def request_export(*, user, document, export_format):
    authorize_feature(user, document.company, Feature.DOCUMENT_EXPORT)
    locked = GeneratedDocument.objects.select_for_update().get(pk=document.pk)
    allowed = (
        GeneratedDocument.Status.APPROVED,
        GeneratedDocument.Status.EXPORTED,
    )
    if locked.status not in allowed:
        raise DocumentStateConflict(
            current_status=locked.status,
            allowed_statuses=allowed,
        )
    revision = _latest_revision(locked)
    if revision is None:
        raise DocumentStateConflict(
            current_status=locked.status,
            allowed_statuses=('approved_with_revision',),
        )
    consume_usage(locked.company, UsageRecord.Metric.DOCUMENT_EXPORT)
    export = DocumentExport.objects.create(
        document=locked,
        revision=revision,
        requested_by=user,
        format=export_format,
        status=DocumentExport.Status.QUEUED,
        metadata={
            'template_id': str(locked.template_id),
            'template_version': locked.template_version,
            'edit_version': locked.edit_version,
            'revision_version': revision.version,
            'context_snapshot': locked.context_snapshot,
        },
    )
    DocumentAuditEvent.objects.create(
        document=locked,
        actor=user,
        event=DocumentAuditEvent.Event.EXPORT_REQUESTED,
        metadata={
            'export_id': str(export.id),
            'format': export_format,
            'revision_version': revision.version,
        },
    )
    return export


@transaction.atomic
def run_export(export_id):
    export = (
        DocumentExport.objects.select_for_update()
        .select_related('document', 'revision')
        .get(pk=export_id)
    )
    if export.status != DocumentExport.Status.QUEUED:
        return export
    export.status = DocumentExport.Status.PROCESSING
    export.started_at = timezone.now()
    export.save(update_fields=['status', 'started_at', 'updated_at'])

    try:
        if export.format == DocumentExport.Format.DOCX:
            payload = _build_docx(export)
            content_type = (
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            )
        else:
            payload = _build_pdf(export)
            content_type = 'application/pdf'
        extension = export.format
        filename = f'{export.document_id}-{export.revision.version}.{extension}'
        export.file.save(filename, ContentFile(payload), save=False)
        export.storage_key = export.file.name
        export.checksum_sha256 = hashlib.sha256(payload).hexdigest()
        export.file_size = len(payload)
        export.status = DocumentExport.Status.COMPLETED
        export.completed_at = timezone.now()
        export.metadata = {
            **export.metadata,
            'content_type': content_type,
        }
        export.error_message = ''
        export.save(
            update_fields=[
                'file',
                'storage_key',
                'checksum_sha256',
                'file_size',
                'status',
                'completed_at',
                'metadata',
                'error_message',
                'updated_at',
            ],
        )
        GeneratedDocument.objects.filter(pk=export.document_id).update(
            status=GeneratedDocument.Status.EXPORTED,
            exported_at=export.completed_at,
            updated_at=export.completed_at,
        )
        DocumentAuditEvent.objects.create(
            document=export.document,
            actor=None,
            event=DocumentAuditEvent.Event.EXPORT_COMPLETED,
            metadata={
                'export_id': str(export.id),
                'format': export.format,
                'checksum_sha256': export.checksum_sha256,
            },
        )
    except Exception as exc:
        logger.exception('Document export failed export=%s', export.id)
        export.status = DocumentExport.Status.FAILED
        export.error_message = 'Document export failed'
        export.failed_at = timezone.now()
        export.save(
            update_fields=[
                'status',
                'error_message',
                'failed_at',
                'updated_at',
            ],
        )
        DocumentAuditEvent.objects.create(
            document=export.document,
            actor=None,
            event=DocumentAuditEvent.Event.EXPORT_FAILED,
            metadata={
                'export_id': str(export.id),
                'format': export.format,
                'error_type': type(exc).__name__,
            },
        )
    return export


def _build_docx(export):
    document = DocxDocument()
    document.core_properties.title = export.revision.title
    document.add_heading(export.revision.title, level=1)
    for line in export.revision.content_text.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_font_name():
    candidates = [
        settings.DOCUMENT_PDF_FONT_PATH,
        'C:/Windows/Fonts/arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            name = 'TenderHelperUnicode'
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, candidate))
            return name
    return 'Helvetica'


def _build_pdf(export):
    buffer = io.BytesIO()
    font_name = _pdf_font_name()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocumentTitle',
        parent=styles['Title'],
        fontName=font_name,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    body_style = ParagraphStyle(
        'DocumentBody',
        parent=styles['BodyText'],
        fontName=font_name,
        leading=16,
        spaceAfter=8,
    )
    story = [Paragraph(_xml_escape(export.revision.title), title_style)]
    for line in export.revision.content_text.splitlines():
        if line.strip():
            story.append(Paragraph(_xml_escape(line.strip()), body_style))
            story.append(Spacer(1, 4))
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=export.revision.title,
        author='TenderHelper',
    )
    pdf.build(story)
    return buffer.getvalue()


def _xml_escape(value):
    return (
        value.replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
    )


def create_download_token(export):
    return TimestampSigner(salt=DOWNLOAD_SALT).sign(str(export.id))


def verify_download_token(export, token):
    try:
        value = TimestampSigner(salt=DOWNLOAD_SALT).unsign(
            token,
            max_age=settings.DOCUMENT_EXPORT_URL_TTL_SECONDS,
        )
    except (BadSignature, SignatureExpired) as exc:
        raise InvalidDownloadToken() from exc
    if value != str(export.id):
        raise InvalidDownloadToken()
    return export
